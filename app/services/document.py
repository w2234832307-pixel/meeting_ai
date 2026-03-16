"""
文档解析服务 - 支持Word和PDF文件
"""
import os
from pathlib import Path
from typing import Optional
from app.core.logger import logger
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentService:
    """文档解析服务类
    
    说明：
    - extract_text_* 系列：只提取纯文本，用于喂给 LLM 做摘要/理解
    - extract_html_* 系列：尽量保留原始格式，生成 HTML5，用于前端展示或格式保留场景
    """
    
    def __init__(self):
        """初始化文档服务"""
        pass
    
    # ============================
    # 纯文本提取接口（原有逻辑）
    # ============================
    def extract_text_from_docx(self, file_path: str) -> Optional[str]:
        """
        从Word文档（.docx）提取文本
        
        Args:
            file_path: Word文档路径
        
        Returns:
            提取的文本内容，失败返回None
        """
        try:
            from docx import Document
            
            logger.info(f"📄 开始解析Word文档: {file_path}")
            
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n".join(paragraphs)
            
            logger.info(f"✅ Word文档解析完成，提取文本长度: {len(full_text)}")
            return full_text
            
        except ImportError:
            logger.error("❌ python-docx 库未安装，请运行: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"❌ Word文档解析失败: {e}")
            return None
    
    def extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        """
        从PDF文档提取文本
        
        Args:
            file_path: PDF文档路径
        
        Returns:
            提取的文本内容，失败返回None
        """
        try:
            import PyPDF2
            
            logger.info(f"📄 开始解析PDF文档: {file_path}")
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 提取所有页面文本
                pages_text = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        pages_text.append(text.strip())
                
                full_text = "\n".join(pages_text)
                
                logger.info(f"✅ PDF文档解析完成，共{len(pdf_reader.pages)}页，提取文本长度: {len(full_text)}")
                return full_text
                
        except ImportError:
            logger.error("❌ PyPDF2 库未安装，请运行: pip install PyPDF2")
            return None
        except Exception as e:
            logger.error(f"❌ PDF文档解析失败: {e}")
            return None
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        自动识别文件类型并提取文本
        
        Args:
            file_path: 文件路径
        
        Returns:
            提取的文本内容，失败返回None
        """
        if not os.path.exists(file_path):
            logger.error(f"❌ 文件不存在: {file_path}")
            return None
        
        # 获取文件扩展名
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.txt':
            # 纯文本文件直接读取
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # 尝试其他编码
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        return f.read()
                except Exception as e:
                    logger.error(f"❌ 文本文件读取失败: {e}")
                    return None
        else:
            logger.error(f"❌ 不支持的文件格式: {ext}")
            return None

    def extract_html_from_docx(self, file_path: str):
        """
        通用的 .docx -> HTML5 转换：
        - 段落：对齐、字号、颜色、粗体、斜体、下划线
        - 标题：根据段落样式名（Heading 1/2/3...）映射为 <h1>/<h2>/<h3>
        - 列表：支持有序/无序列表（基于段落样式名推断）
        - 表格：输出 <table><tr><td>...</td></tr></table>
        
        注意：不会依赖某个具体模板的内容，所有逻辑基于通用的 Word 结构信息。
        """
        from docx import Document
        from docx.enum.text import WD_UNDERLINE

        try:
            doc = Document(file_path)
            html_parts = []

            # 1. 先处理表格（保持顺序：段落和表格在 doc._body._element 中混合比较复杂，这里简化为“先段落后表格”）
            #   如果将来有强需求，可以按 document._body._element 遍历以完全还原顺序。

            # 1.1 处理普通段落（含标题/列表）
            html_parts.append(self._convert_paragraphs_to_html(doc.paragraphs))

            # 1.2 处理表格
            if doc.tables:
                for table in doc.tables:
                    html_parts.append(self._convert_table_to_html(table))

            return "".join(html_parts)

        except Exception as e:
            import traceback
            return f"解析异常: {str(e)}\n{traceback.format_exc()}"

    # ============================
    # 内部辅助方法：段落 / 列表 / 表格
    # ============================

    def _convert_paragraphs_to_html(self, paragraphs):
        """将段落列表转换为 HTML，支持标题、列表、普通段落。"""
        html_parts = []
        i = 0
        n = len(paragraphs)

        while i < n:
            p = paragraphs[i]
            text = p.text or ""
            style_name = (p.style.name or "").lower() if p.style else ""

            # 1. 检测列表段落（基于样式名）
            is_bullet = any(k in style_name for k in ["bullet", "list bullet", "项目符号"])
            is_number = any(k in style_name for k in ["number", "list number", "编号"])

            if (is_bullet or is_number) and text.strip():
                # 聚合连续的列表段落
                list_items = []
                list_type = "ul" if is_bullet else "ol"

                while i < n:
                    p2 = paragraphs[i]
                    style2 = (p2.style.name or "").lower() if p2.style else ""
                    is_bullet2 = any(k in style2 for k in ["bullet", "list bullet", "项目符号"])
                    is_number2 = any(k in style2 for k in ["number", "list number", "编号"])
                    if not (is_bullet2 or is_number2):
                        break
                    if (p2.text or "").strip():
                        item_html = self._convert_single_paragraph_runs_to_html(p2)
                        list_items.append(f"<li>{item_html}</li>")
                    i += 1

                if list_items:
                    html_parts.append(f"<{list_type}>" + "".join(list_items) + f"</{list_type}>")
                continue  # continue while

            # 2. 非列表：判断是否为标题
            heading_level = None
            if style_name.startswith("heading"):
                # e.g. "Heading 1" -> 1
                parts = style_name.split()
                for p_token in parts:
                    if p_token.isdigit():
                        heading_level = int(p_token)
                        break

            # 3. 普通段落/标题
            runs_html, align_class, inline_style = self._convert_paragraph_style_and_runs(p)

            if runs_html.strip():
                if heading_level and 1 <= heading_level <= 6:
                    tag = f"h{heading_level}"
                else:
                    tag = "p"

                class_attr = f' class="{align_class}"' if align_class else ""
                style_attr = f' style="{inline_style}"' if inline_style else ""
                html_parts.append(f"<{tag}{class_attr}{style_attr}>{runs_html}</{tag}>")

            i += 1

        return "".join(html_parts)

    def _convert_paragraph_style_and_runs(self, p):
        """将单个段落的对齐 / 红头线 / run 样式转换为 HTML 片段。"""
        # 1. 对齐检测（与原逻辑一致）
        alignment = p.alignment
        if alignment is None:
            alignment = p.paragraph_format.alignment
        if alignment is None and p.style:
            curr_style = p.style
            while curr_style:
                if curr_style.paragraph_format.alignment is not None:
                    alignment = curr_style.paragraph_format.alignment
                    break
                curr_style = getattr(curr_style, "base_style", None)

        # 2. 大字号兜底：推断标题居中
        is_large_font = False
        for run in p.runs:
            if run.font and run.font.size and getattr(run.font.size, "pt", None):
                if run.font.size.pt > 20:
                    is_large_font = True
                    break

        align_class = ""
        if alignment == 1 or is_large_font:  # 1 = center
            align_class = "center"
        elif alignment == 2:  # 2 = right
            align_class = "right"

        # 3. 红头横线逻辑（保持兼容）
        line_style = ""
        if "编号" in p.text or "签发人" in p.text:
            line_style = "border-bottom: 2px solid black; padding-bottom: 5px; margin-bottom: 15px;"

        # 4. 汇总 run 样式
        runs_html = ""
        for run in p.runs:
            runs_html += self._convert_run_to_html(run)

        # 5. 计算段落的 inline style（对齐 + 红头线）
        inline_style = line_style or ""
        if align_class == "center":
            inline_style += " text-align: center !important;"
        elif align_class == "right":
            inline_style += " text-align: right !important;"
        inline_style = inline_style.strip()

        return runs_html, align_class, inline_style

    def _convert_run_to_html(self, run):
        """将单个 run 转为带样式的 <span>/<strong> 结构，支持粗体、斜体、下划线、颜色、字号。"""
        text = run.text or ""
        if text == "":
            return ""

        # 基础字体样式
        font_css = []
        if run.font is not None and getattr(run.font, "size", None):
            try:
                pt = run.font.size.pt
                if pt:
                    font_css.append(f"font-size: {pt * 1.3}px")
            except Exception:
                pass
        if run.font is not None and getattr(run.font, "color", None) and getattr(run.font.color, "rgb", None):
            font_css.append(f"color: #{run.font.color.rgb}")

        style_str = "; ".join(font_css)
        content = text.replace(" ", "&nbsp;")  # 保留空格

        # 处理下划线和斜体
        extra_style = []
        if getattr(run, "underline", False):
            extra_style.append("text-decoration: underline;")
        if getattr(run, "italic", False):
            extra_style.append("font-style: italic;")
        if extra_style:
            if style_str:
                style_str = style_str + "; " + " ".join(extra_style)
            else:
                style_str = " ".join(extra_style)

        inner_html = f'<span style="{style_str}">{content}</span>'

        # 加粗
        if getattr(run, "bold", False):
            inner_html = f'<strong style="font-weight: bold;">{inner_html}</strong>'

        return inner_html

    def _convert_table_to_html(self, table):
        """将 Word 表格转换为 HTML <table> 结构（基础结构 + 粗体表头）。"""
        rows = table.rows
        if not rows:
            return ""

        html_rows = []
        for r_idx, row in enumerate(rows):
            cell_htmls = []
            for cell in row.cells:
                # 合并所有段落的 HTML
                cell_text = self._convert_paragraphs_to_html(cell.paragraphs)
                tag = "th" if r_idx == 0 else "td"
                cell_htmls.append(f"<{tag}>{cell_text}</{tag}>")
            html_rows.append("<tr>" + "".join(cell_htmls) + "</tr>")

        return "<table>" + "".join(html_rows) + "</table>"


    def extract_html_from_pdf(self, file_path: str) -> Optional[str]:
        """
        从 PDF 文档提取带格式的 HTML。
        
        说明：
        - 高保真还原依赖外部工具 pdf2htmlEX（推荐），需要系统已安装该命令。
        - 如果 pdf2htmlEX 不可用，则退化为简单的文本包装成 <pre> 的 HTML。
        """
        import subprocess
        import tempfile

        logger.info(f"📄 开始将 PDF 文档转换为 HTML: {file_path}")

        # 先尝试使用 pdf2htmlEX（需要系统安装）
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as tmp_html:
                output_path = tmp_html.name

            cmd = [
                "pdf2htmlEX",
                "--embed", "cfijo",  # 内嵌字体/图片/JS/CSS，方便单文件传输
                "--dest-dir", os.path.dirname(output_path),
                "--output", os.path.basename(output_path),
                file_path,
            ]
            logger.info(f"🔧 调用 pdf2htmlEX: {' '.join(cmd)}")
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

            if os.path.exists(output_path):
                with open(output_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_content = f.read()
                logger.info(f"✅ PDF->HTML 转换完成（pdf2htmlEX），长度: {len(html_content)}")
                return html_content
        except FileNotFoundError:
            logger.warning("⚠️ 未找到 pdf2htmlEX 命令，将退化为简单的文本 HTML。请在系统中安装 pdf2htmlEX 以获得高保真还原。")
        except subprocess.CalledProcessError as e:
            logger.warning(f"⚠️ pdf2htmlEX 转换失败，退化为简单文本 HTML: {e}")
        except Exception as e:
            logger.warning(f"⚠️ pdf2htmlEX 调用异常，退化为简单文本 HTML: {e}")

        # 退化方案：使用纯文本 + <pre> 包裹（不会高保真，但至少可用）
        text = self.extract_text_from_pdf(file_path)
        if text is None:
            return None
        safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        fallback_html = (
            "<!DOCTYPE html>\n"
            "<html><head><meta charset=\"utf-8\"></head><body>\n"
            "<pre style=\"white-space: pre-wrap;\">\n"
            f"{safe_text}\n"
            "</pre>\n"
            "</body></html>"
        )
        logger.info(f"✅ 使用退化方案生成 PDF 文本 HTML，长度: {len(fallback_html)}")
        return fallback_html

    def extract_html_from_file(self, file_path: str) -> Optional[str]:
        """
        自动识别文件类型并提取 HTML（尽量保留原始格式）
        
        - .docx -> 语义化 HTML（mammoth）
        - .pdf  -> 优先 pdf2htmlEX，高保真 HTML；否则退化为 <pre> 文本 HTML
        - .txt  -> 简单包裹到 <pre> 中
        """
        if not os.path.exists(file_path):
            logger.error(f"❌ 文件不存在: {file_path}")
            return None

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            return self.extract_html_from_docx(file_path)
        elif ext == ".pdf":
            return self.extract_html_from_pdf(file_path)
        elif ext == ".txt":
            # 纯文本文件：简单用 <pre> 包装
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except UnicodeDecodeError:
                try:
                    with open(file_path, "r", encoding="gbk") as f:
                        text = f.read()
                except Exception as e:
                    logger.error(f"❌ 文本文件读取失败: {e}")
                    return None
            safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html = (
                "<!DOCTYPE html>\n"
                "<html><head><meta charset=\"utf-8\"></head><body>\n"
                "<pre style=\"white-space: pre-wrap;\">\n"
                f"{safe_text}\n"
                "</pre>\n"
                "</body></html>"
            )
            logger.info(f"✅ 文本文件转 HTML 完成，长度: {len(html)}")
            return html
        else:
            logger.error(f"❌ 不支持的文件格式(HTML提取): {ext}")
            return None


# 创建单例实例
document_service = DocumentService()